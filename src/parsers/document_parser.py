"""Document parser for extracting text from PDF and DOCX files"""

import re
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

try:
    import fitz  # PyMuPDF
except ImportError:
    # Fallback for PyMuPDF4llm
    try:
        import pymupdf4llm
        fitz = None
    except ImportError:
        fitz = None

import docx2txt
from docx import Document


@dataclass
class ParsedDocument:
    """Container for parsed document content"""
    text: str
    filename: str
    file_type: str
    metadata: dict


class DocumentParser:
    """Parse PDF and DOCX documents into clean text"""
    
    def __init__(self):
        self.supported_formats = {".pdf", ".docx", ".doc"}
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file and extract text
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with extracted text and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = path.suffix.lower()
        if suffix not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {suffix}. Supported: {self.supported_formats}")
        
        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix in {".docx", ".doc"}:
            return self._parse_docx(path)
    
    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Extract text from PDF using PyMuPDF or fallback"""
        try:
            if fitz:
                # Use PyMuPDF if available
                doc = fitz.open(path)
                text_parts = []
                
                for page_num, page in enumerate(doc):
                    text = page.get_text()
                    text_parts.append(text)
                
                doc.close()
                
                full_text = "\n".join(text_parts)
                cleaned_text = self._clean_text(full_text)
                
                return ParsedDocument(
                    text=cleaned_text,
                    filename=path.name,
                    file_type="pdf",
                    metadata={"pages": len(text_parts)}
                )
            else:
                # Fallback to pdfminer
                from pdfminer.high_level import extract_text
                text = extract_text(str(path))
                cleaned_text = self._clean_text(text)
                
                return ParsedDocument(
                    text=cleaned_text,
                    filename=path.name,
                    file_type="pdf",
                    metadata={"pages": 1}
                )
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {e}")
    
    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Extract text from DOCX"""
        try:
            # Try python-docx first for better structure
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            
            # Fallback to docx2txt if needed
            if not text.strip():
                text = docx2txt.process(str(path))
            
            cleaned_text = self._clean_text(text)
            
            return ParsedDocument(
                text=cleaned_text,
                filename=path.name,
                file_type="docx",
                metadata={"paragraphs": len(paragraphs)}
            )
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page numbers and common artifacts
        text = re.sub(r'\n\d+\n', '\n', text)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        return text.strip()
    
    def parse_text(self, text: str, filename: str = "text_input") -> ParsedDocument:
        """Parse raw text input (for JD pasted as text)"""
        cleaned_text = self._clean_text(text)
        
        return ParsedDocument(
            text=cleaned_text,
            filename=filename,
            file_type="text",
            metadata={}
        )
